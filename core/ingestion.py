"""Document ingestion and chunking utilities."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

import fitz
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


@dataclass(frozen=True)
class SourceChunk:
    """Lightweight representation of an extracted text chunk."""

    text: str
    metadata: dict[str, object]


def is_supported_file(file_path: str | Path) -> bool:
    """Return True when the file extension is one we can ingest."""

    return Path(file_path).suffix.lower() in SUPPORTED_EXTENSIONS


def load_documents(file_paths: Sequence[str | Path]) -> list[Document]:
    """Load and split multiple files into LangChain documents."""

    chunks: list[Document] = []
    for file_path in file_paths:
        chunks.extend(ingest_file(file_path))
    return chunks


def ingest_file(file_path: str | Path) -> list[Document]:
    """Load a supported file and split it into metadata-rich chunks."""

    path = Path(file_path)
    if not is_supported_file(path):
        raise ValueError(f"Unsupported file type: {path.suffix}")

    raw_chunks = _extract_source_chunks(path)
    documents: list[Document] = []

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    for raw_chunk in raw_chunks:
        split_texts = splitter.split_text(raw_chunk.text)
        if not split_texts:
            continue

        for index, split_text in enumerate(split_texts, start=1):
            metadata = dict(raw_chunk.metadata)
            metadata["chunk_index"] = index
            metadata["chunk_count"] = len(split_texts)
            documents.append(Document(page_content=split_text, metadata=metadata))

    return documents


def _extract_source_chunks(path: Path) -> list[SourceChunk]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_chunks(path)
    if suffix == ".docx":
        return _extract_docx_chunks(path)
    if suffix == ".txt":
        return _extract_txt_chunks(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf_chunks(path: Path) -> list[SourceChunk]:
    chunks: list[SourceChunk] = []

    with fitz.open(path) as pdf:
        for page_number, page in enumerate(pdf, start=1):
            text = page.get_text("text").strip()
            if not text:
                continue

            metadata = {
                "source": path.name,
                "file_path": str(path),
                "page": page_number,
                "section": _guess_section_from_text(text),
                "source_type": "pdf",
            }
            chunks.append(SourceChunk(text=text, metadata=metadata))

    return chunks


def _extract_docx_chunks(path: Path) -> list[SourceChunk]:
    doc = DocxDocument(path)
    chunks: list[SourceChunk] = []
    current_section = "Document"
    paragraph_buffer: list[str] = []

    def flush_buffer(section_name: str) -> None:
        nonlocal paragraph_buffer
        text = "\n".join(line for line in paragraph_buffer if line.strip()).strip()
        if text:
            chunks.append(
                SourceChunk(
                    text=text,
                    metadata={
                        "source": path.name,
                        "file_path": str(path),
                        "page": None,
                        "section": section_name,
                        "source_type": "docx",
                    },
                )
            )
        paragraph_buffer = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.lower().startswith("heading"):
            flush_buffer(current_section)
            current_section = text
            continue

        paragraph_buffer.append(text)

    flush_buffer(current_section)
    return chunks


def _extract_txt_chunks(path: Path) -> list[SourceChunk]:
    text = path.read_text(encoding="utf-8").strip()
    if not text:
        return []

    return [
        SourceChunk(
            text=text,
            metadata={
                "source": path.name,
                "file_path": str(path),
                "page": None,
                "section": "Document",
                "source_type": "txt",
            },
        )
    ]


def _guess_section_from_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return "Document"

    first_line = lines[0]
    if len(first_line) <= 120 and not first_line.endswith("."):
        return first_line
    return "Document"


def collect_supported_files(paths: Iterable[str | Path]) -> list[Path]:
    """Filter and normalize a set of file paths to supported inputs."""

    supported_paths: list[Path] = []
    for file_path in paths:
        path = Path(file_path)
        if path.is_file() and is_supported_file(path):
            supported_paths.append(path)
    return supported_paths
